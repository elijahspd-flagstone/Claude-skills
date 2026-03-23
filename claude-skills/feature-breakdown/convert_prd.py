#!/usr/bin/env python3
"""
PRD Converter
Converts a PRD markdown file to branded PDF and DOCX formats.

Usage:
    python3 convert_prd.py path/to/prd.md
    python3 convert_prd.py path/to/prd.md --output ./exports
    python3 convert_prd.py path/to/prd.md --pdf-only
    python3 convert_prd.py path/to/prd.md --docx-only

Dependencies:
    pip install python-docx reportlab
"""

import argparse
import re
import sys
from pathlib import Path

# ── Brand colors — imported from single source of truth ──────────────────────
from brand_colors import (
    NAVY, GOLD, ICE_BLUE, OFF_WHITE, SLATE, GRAY, WARM_TAN,
)


# ── Skill config — read skill-config.md from the same directory ───────────────

def _load_config() -> dict:
    """Parse skill-config.md (key: value pairs) from the skill directory."""
    cfg_path = Path(__file__).parent / "skill-config.md"
    result = {}
    if not cfg_path.exists():
        return result
    for line in cfg_path.read_text(encoding="utf-8").splitlines():
        m = re.match(r"^(\w+):\s+(.+)$", line.strip())
        if m:
            result[m.group(1)] = m.group(2).strip()
    return result


_cfg = _load_config()

COMPANY_NAME = _cfg.get("company_name", "Codepup")

_logo_str = _cfg.get("logo_path", "")
if _logo_str:
    LOGO_PATH = Path(_logo_str)
else:
    LOGO_PATH = Path(__file__).parent / "assets" / "codepup_logo.png"

WATERMARK_OPACITY = 0.02   # 2% — barely there


def hex_to_rgb(hex_color: str):
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


# ── Shared markdown pre-processor ────────────────────────────────────────────

def preprocess(md_text: str) -> list[str]:
    """
    Unwrap any ```markdown / ````markdown fenced blocks by replacing them with
    their inner content (rendered as markdown, not as code).  All other fenced
    code blocks are left intact.  Returns a flat list of lines.
    """
    lines = md_text.split("\n")
    out = []
    i = 0
    while i < len(lines):
        m = re.match(r"^(`{3,})\s*(markdown)?\s*$", lines[i].rstrip(), re.IGNORECASE)
        if m:
            fence = m.group(1)          # the opening fence (e.g. ```` or ```)
            lang  = (m.group(2) or "").lower()
            # Collect until closing fence of same length
            i += 1
            block = []
            while i < len(lines):
                close = re.match(r"^(`{" + str(len(fence)) + r",})\s*$", lines[i].rstrip())
                if close:
                    i += 1
                    break
                block.append(lines[i])
                i += 1
            if lang == "markdown":
                # Inline-expand: recurse so nested fences also get handled
                out.extend(preprocess("\n".join(block)))
            else:
                out.append(fence)           # re-emit the opening fence
                out.extend(block)
                out.append(fence)           # re-emit the closing fence
        else:
            out.append(lines[i])
            i += 1
    return out


# ── PDF via reportlab ─────────────────────────────────────────────────────────

def _render_mermaid(mmd_content: str, out_path: Path) -> bool:
    """
    Render mermaid syntax to a PNG at out_path using npx @mermaid-js/mermaid-cli.
    Theme colors are sourced from brand_colors.py — no duplicated hex values.
    Returns True on success.
    """
    import subprocess, tempfile, json
    from brand_colors import (
        NAVY, GOLD, ICE_BLUE, OFF_WHITE, SLATE, GRAY, WARM_TAN,
    )
    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)

        # Build mermaid theme config from brand colors
        mermaid_config = {
            "theme": "base",
            "themeVariables": {
                "primaryColor":        OFF_WHITE,
                "primaryTextColor":    NAVY,
                "primaryBorderColor":  NAVY,
                "lineColor":           SLATE,
                "secondaryColor":      ICE_BLUE,
                "tertiaryColor":       OFF_WHITE,
                "edgeLabelBackground": "#FFFFFF",
                "clusterBkg":          ICE_BLUE,
                "titleColor":          NAVY,
                "fontFamily":          "Helvetica, Arial, sans-serif",
                "fontSize":            "15px",
            },
        }

        with tempfile.NamedTemporaryFile(suffix=".mmd", mode="w",
                                         delete=False, encoding="utf-8") as f:
            f.write(mmd_content)
            mmd_path = Path(f.name)

        with tempfile.NamedTemporaryFile(suffix=".json", mode="w",
                                         delete=False, encoding="utf-8") as cf:
            json.dump(mermaid_config, cf)
            cfg_path = Path(cf.name)

        result = subprocess.run(
            ["npx", "-y", "@mermaid-js/mermaid-cli",
             "-i", str(mmd_path), "-o", str(out_path),
             "--configFile", str(cfg_path),
             "--width", "1400", "--backgroundColor", "transparent"],
            capture_output=True, timeout=120,
        )
        mmd_path.unlink(missing_ok=True)
        cfg_path.unlink(missing_ok=True)
        return result.returncode == 0 and out_path.exists()
    except Exception:
        return False


def _make_watermark_image(opacity: float) -> "Path | None":
    """Return a temp PNG with the logo at the given opacity (0–1), or None."""
    if not LOGO_PATH.exists():
        return None
    try:
        from PIL import Image
        import io, tempfile
        img = Image.open(LOGO_PATH).convert("RGBA")
        r, g, b, a = img.split()
        a = a.point(lambda v: int(v * opacity))
        img.putalpha(a)
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        img.save(tmp.name, "PNG")
        return Path(tmp.name)
    except Exception:
        return None


def convert_to_pdf(md_path: Path, output_dir: Path) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, Preformatted,
        ListFlowable, ListItem,
    )
    from reportlab.platypus.flowables import HRFlowable

    md_text = md_path.read_text(encoding="utf-8")
    title = md_path.stem.replace("-", " ").replace("_", " ").title()
    out_path = output_dir / (md_path.stem + ".pdf")

    watermark_img = _make_watermark_image(WATERMARK_OPACITY)
    page_w, page_h = A4

    def on_page(canvas, doc):
        canvas.saveState()

        # ── Watermark ──
        if watermark_img and watermark_img.exists():
            wm_size = page_w * 0.68
            canvas.drawImage(
                str(watermark_img),
                (page_w - wm_size) / 2, (page_h - wm_size) / 2,
                width=wm_size, height=wm_size,
                preserveAspectRatio=True, mask="auto",
            )

        # ── Running page header ──
        hdr_y = page_h - 1.6 * cm
        canvas.setFont("Helvetica-Bold", 12)
        canvas.setFillColor(colors.HexColor(NAVY))
        canvas.drawString(2.2 * cm, hdr_y, COMPANY_NAME)
        name_w = canvas.stringWidth(COMPANY_NAME, "Helvetica-Bold", 12)
        canvas.setFont("Helvetica", 10)
        canvas.setFillColor(colors.HexColor(WARM_TAN))
        canvas.drawString(2.2 * cm + name_w + 6, hdr_y,
                          "— Product Requirements Document")
        # Gold rule under header
        canvas.setStrokeColor(colors.HexColor(GOLD))
        canvas.setLineWidth(1.2)
        rule_y = hdr_y - 5
        canvas.line(2.2 * cm, rule_y, page_w - 2.2 * cm, rule_y)

        # ── Running page footer ──
        ftr_y = 1.4 * cm
        canvas.setStrokeColor(colors.HexColor(GOLD))
        canvas.setLineWidth(0.8)
        canvas.line(2.2 * cm, ftr_y + 8, page_w - 2.2 * cm, ftr_y + 8)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(WARM_TAN))
        canvas.drawString(2.2 * cm, ftr_y, f"{COMPANY_NAME} — Confidential")
        page_label = f"Page {doc.page}"
        lbl_w = canvas.stringWidth(page_label, "Helvetica", 7.5)
        canvas.drawString(page_w - 2.2 * cm - lbl_w, ftr_y, page_label)

        canvas.restoreState()

    c_navy     = colors.HexColor(NAVY)
    c_gold     = colors.HexColor(GOLD)
    c_ice      = colors.HexColor(ICE_BLUE)
    c_slate    = colors.HexColor(SLATE)
    c_gray     = colors.HexColor(GRAY)
    c_tan      = colors.HexColor(WARM_TAN)
    c_offwhite = colors.HexColor(OFF_WHITE)

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sty = {
        "h1":    S("h1",   fontName="Helvetica-Bold", fontSize=20, textColor=c_navy,
                            leading=26, spaceBefore=4,  spaceAfter=8),
        "h2":    S("h2",   fontName="Helvetica-Bold", fontSize=16, textColor=c_navy,
                            leading=20, spaceBefore=20, spaceAfter=6),
        "h3":    S("h3",   fontName="Helvetica-Bold", fontSize=12, textColor=c_slate,
                            leading=16, spaceBefore=14, spaceAfter=4),
        "h4":    S("h4",   fontName="Helvetica-Bold", fontSize=10, textColor=c_slate,
                            leading=13, spaceBefore=8,  spaceAfter=3),
        "body":  S("body", fontName="Helvetica",      fontSize=10, textColor=c_slate,
                            leading=15, spaceBefore=2,  spaceAfter=5),
        "li":    S("li",   fontName="Helvetica",      fontSize=10, textColor=c_slate,
                            leading=14, spaceBefore=1,  spaceAfter=2, leftIndent=14),
        "li2":   S("li2",  fontName="Helvetica",      fontSize=10, textColor=c_slate,
                            leading=14, spaceBefore=1,  spaceAfter=2, leftIndent=28),
        "code":  S("code", fontName="Courier",        fontSize=8.5, textColor=c_navy,
                            leading=12, spaceBefore=2,  spaceAfter=2,
                            backColor=c_offwhite, leftIndent=8, rightIndent=8),
        "quote": S("quote",fontName="Helvetica",      fontSize=9.5, textColor=c_navy,
                            leading=14, spaceBefore=4,  spaceAfter=6,
                            leftIndent=12, backColor=c_ice),
        "footer":S("footer",fontName="Helvetica",     fontSize=8,  textColor=c_tan,
                            leading=11),
        "brand": S("brand", fontName="Helvetica-Bold",fontSize=13, textColor=c_navy,
                            leading=16, spaceAfter=2),
        "meta":  S("meta",  fontName="Helvetica-Oblique", fontSize=9, textColor=c_tan,
                            leading=14, spaceBefore=1, spaceAfter=1,
                            backColor=c_offwhite, leftIndent=6, rightIndent=6),
    }

    def hr(color=None, thickness=1):
        return HRFlowable(width="100%", thickness=thickness,
                          color=color or c_gray, spaceAfter=4, spaceBefore=4)

    def inline_fmt(text: str) -> str:
        text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", text)
        text = re.sub(r"\*\*(.+?)\*\*",     r"<b>\1</b>",         text)
        text = re.sub(r"(?<!\*)\*(.+?)\*(?!\*)", r"<i>\1</i>",    text)
        text = re.sub(r"`(.+?)`", r'<font name="Courier" size="8.5">\1</font>', text)
        text = re.sub(r"&(?!#?\w+;)", "&amp;", text)
        return text

    story         = []
    diagram_count = 0
    story.append(Spacer(1, 4))   # small top gap below running header

    # ── List buffer — collect consecutive list items, flush as ListFlowable ──
    # Each entry: ("bullet"|"check"|"ordered"|"sub", text, number)
    list_buf = []

    # Styles for check/sub rendered as plain paragraphs
    sty_check = S("check", fontName="Helvetica", fontSize=10, textColor=c_slate,
                  leading=15, spaceBefore=1, spaceAfter=2, leftIndent=20, firstLineIndent=-14)
    sty_sub   = S("sub",   fontName="Helvetica", fontSize=9.5, textColor=c_slate,
                  leading=14, spaceBefore=1, spaceAfter=2, leftIndent=36)

    # Hanging-indent style for ordered list items — number sits in the overhang
    sty_ordered = S("ordered", fontName="Helvetica", fontSize=10, textColor=c_slate,
                    leading=14, spaceBefore=2, spaceAfter=2,
                    leftIndent=36, firstLineIndent=-24)

    def _flush_flow_items(flow_buf):
        """Render a batch of bullet/ordered items."""
        if not flow_buf:
            return
        has_ordered = any(k == "ordered" for k, _, _ in flow_buf)
        if has_ordered:
            # Render ordered items as hanging-indent paragraphs for precise number placement
            for kind, text, num in flow_buf:
                label = f"<b>{num}.</b>&nbsp;&nbsp;" if num is not None else "•&nbsp;&nbsp;"
                story.append(Paragraph(label + inline_fmt(text), sty_ordered))
        else:
            items = []
            for kind, text, num in flow_buf:
                p = Paragraph(inline_fmt(text), sty["body"])
                li = ListItem(p, leftIndent=22,
                              bulletColor=c_navy, bulletFontSize=8)
                items.append(li)
            lf = ListFlowable(
                items,
                bulletType="bullet",
                bulletFontName="Helvetica",
                bulletFontSize=8,
                bulletColor=c_navy,
                leftIndent=8,
                spaceBefore=2,
                spaceAfter=6,
            )
            story.append(lf)

    def flush_list():
        if not list_buf:
            return
        flow_buf = []   # accumulates bullet/ordered items for ListFlowable
        for kind, text, num in list_buf:
            if kind in ("bullet", "ordered", "check"):
                # check items render identically to bullets — consistent dot style
                flow_buf.append(("bullet" if kind == "check" else kind, text, num))
            else:
                _flush_flow_items(flow_buf)
                flow_buf = []
                # sub items (Acceptance lines) — indented paragraph with dash
                story.append(Paragraph(
                    '–  ' + inline_fmt(text), sty_sub))
        _flush_flow_items(flow_buf)
        list_buf.clear()

    # ── Parse pre-processed lines ──
    lines = preprocess(md_text)
    i = 0
    in_code    = False
    code_buf   = []
    table_rows = []

    def flush_table():
        if not table_rows:
            return []
        col_count = max(len(r) for r in table_rows)
        padded    = [r + [""] * (col_count - len(r)) for r in table_rows]
        col_width = (A4[0] - 2.2 * cm - 2.2 * cm) / col_count
        tdata = [[Paragraph(inline_fmt(c.strip()), sty["body"]) for c in row]
                 for row in padded]
        ts = TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  c_ice),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  c_navy),
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("GRID",          (0, 0), (-1, -1), 0.5, c_gray),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, c_offwhite]),
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ])
        t = Table(tdata, colWidths=[col_width] * col_count)
        t.setStyle(ts)
        return [t, Spacer(1, 6)]

    def is_list_line(ln):
        return (re.match(r"^- \[[ xX]\] ", ln) or
                re.match(r"^  [*\-] ",     ln) or
                re.match(r"^[*\-] ",       ln) or
                re.match(r"^\d+\. ",       ln))

    ord_counter = 0

    while i < len(lines):
        line = lines[i]

        # Flush list buffer when a non-list line is about to be processed
        if list_buf and not is_list_line(line):
            flush_list()

        # Fenced code blocks
        m_fence = re.match(r"^(`{3,})\s*(\S*)\s*$", line.rstrip())
        if m_fence and not in_code:
            in_code    = True
            code_lang  = (m_fence.group(2) or "").lower()
            code_buf   = []
            i += 1
            continue
        if in_code:
            if re.match(r"^(`{3,})\s*$", line.rstrip()):
                in_code = False
                if code_lang == "mermaid":
                    diagram_count += 1
                    diagrams_dir = md_path.parent / "diagrams"
                    png_path = diagrams_dir / f"{md_path.stem}-diagram-{diagram_count}.png"
                    ok = _render_mermaid("\n".join(code_buf), png_path)
                    if ok:
                        from reportlab.platypus import Image as RLImage
                        available_w = page_w - 2 * 2.2 * cm
                        img = RLImage(str(png_path), width=available_w,
                                      height=available_w * 0.6,
                                      kind="proportional")
                        story.append(Spacer(1, 6))
                        story.append(img)
                        story.append(Spacer(1, 8))
                    else:
                        # Fallback: render as plain code if mmdc fails
                        story.append(Preformatted("\n".join(code_buf), sty["code"]))
                        story.append(Spacer(1, 4))
                else:
                    story.append(Preformatted("\n".join(code_buf), sty["code"]))
                    story.append(Spacer(1, 4))
            else:
                code_buf.append(line)
            i += 1
            continue

        # Tables
        if "|" in line and line.strip().startswith("|"):
            cells  = [c for c in line.strip().split("|") if c.strip()]
            is_sep = all(re.match(r"[-: ]+$", c.strip()) for c in cells)
            if not is_sep:
                table_rows.append(cells)
            i += 1
            if i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                continue
            story.extend(flush_table())
            table_rows = []
            continue

        if table_rows:
            story.extend(flush_table())
            table_rows = []

        # Headings
        if   line.startswith("# "):
            ord_counter = 0
            story.append(Paragraph(inline_fmt(line[2:].strip()), sty["h1"]))
        elif line.startswith("## "):
            ord_counter = 0
            story.append(hr(color=c_gray, thickness=0.75))
            story.append(Paragraph(inline_fmt(line[3:].strip()), sty["h2"]))
        elif line.startswith("### "):
            ord_counter = 0
            story.append(Paragraph(inline_fmt(line[4:].strip()), sty["h3"]))
        elif line.startswith("#### "):
            story.append(Paragraph(inline_fmt(line[5:].strip()), sty["h4"]))
        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            story.append(hr())
        # Blockquote
        elif line.startswith("> "):
            story.append(Paragraph(inline_fmt(line[2:]), sty["quote"]))
        # Checklist
        elif re.match(r"^- \[[ xX]\] ", line):
            text = re.sub(r"^- \[[ xX]\] ", "", line)
            list_buf.append(("check", text, None))
        # Indented sub-bullet
        elif re.match(r"^  [*\-] ", line):
            list_buf.append(("sub", line[4:].strip(), None))
        # Unordered list
        elif re.match(r"^[*\-] ", line):
            list_buf.append(("bullet", line[2:].strip(), None))
        # Ordered list
        elif re.match(r"^\d+\. ", line):
            m = re.match(r"^(\d+)\. (.+)", line)
            if m:
                ord_counter += 1
                list_buf.append(("ordered", m.group(2), ord_counter))
        # Empty line
        elif line.strip() == "":
            story.append(Spacer(1, 3))
        # Meta info lines (PRD stamp block at top of document)
        elif (line.startswith("PRD generated by") or
              re.match(r"^Feature:", line) or
              re.match(r"^Status:", line) or
              re.match(r"^Author:", line) or
              re.match(r"^Date:", line)):
            story.append(Paragraph(inline_fmt(line), sty["meta"]))
        # Normal paragraph
        else:
            ord_counter = 0
            story.append(Paragraph(inline_fmt(line), sty["body"]))

        i += 1

    flush_list()  # flush any trailing list

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=3.5*cm, rightMargin=3.5*cm,
        topMargin=2.6*cm,  bottomMargin=2.4*cm,
        title=f"{title} — {COMPANY_NAME} PRD", author=COMPANY_NAME,
    )
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)

    if watermark_img and watermark_img.exists():
        watermark_img.unlink(missing_ok=True)

    return out_path


# ── DOCX via python-docx ──────────────────────────────────────────────────────

def _add_docx_watermark(doc) -> None:
    """
    Adds the company logo as a centred, behind-text watermark to every page
    by inserting a picture into each section's header with Word's watermark XML.
    """
    if not LOGO_PATH.exists():
        return
    try:
        from PIL import Image as PILImage
        import io as _io, tempfile as _tmp
        img = PILImage.open(LOGO_PATH).convert("RGBA")
        r, g, b, a = img.split()
        a = a.point(lambda v: int(v * WATERMARK_OPACITY))
        img.putalpha(a)
        buf = _io.BytesIO()
        img.save(buf, "PNG")
        buf.seek(0)

        from docx.shared import Inches, Pt, Emu
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        WMK_INCHES = 5.5   # watermark width in inches — fills A4 page nicely

        for section in doc.sections:
            header = section.header
            # Clear any existing paragraphs
            for p in header.paragraphs:
                p.clear()
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)

            run = para.add_run()
            # Add the image as an inline picture first (so python-docx handles the rId)
            pic = run.add_picture(buf, width=Inches(WMK_INCHES))
            buf.seek(0)  # rewind for next section

            # Now grab the <wp:inline> element and convert it to <wp:anchor>
            # so the image floats behind the text, centred on the page.
            inline = run._r.find(qn("wp:inline"))
            if inline is None:
                continue

            # Build <wp:anchor> with behind-text positioning
            anchor = OxmlElement("wp:anchor")
            anchor.set("distT", "0"); anchor.set("distB", "0")
            anchor.set("distL", "0"); anchor.set("distR", "0")
            anchor.set("simplePos", "0")
            anchor.set("relativeHeight", "251658240")   # z-order: behind text
            anchor.set("behindDoc",      "1")
            anchor.set("locked",         "0")
            anchor.set("layoutInCell",   "1")
            anchor.set("allowOverlap",   "1")

            # <wp:simplePos>
            sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0")
            anchor.append(sp)

            # <wp:positionH> — centred horizontally on the page
            ph = OxmlElement("wp:positionH"); ph.set("relativeFrom", "page")
            ph_align = OxmlElement("wp:align"); ph_align.text = "center"
            ph.append(ph_align); anchor.append(ph)

            # <wp:positionV> — centred vertically on the page
            pv = OxmlElement("wp:positionV"); pv.set("relativeFrom", "page")
            pv_align = OxmlElement("wp:align"); pv_align.text = "center"
            pv.append(pv_align); anchor.append(pv)

            # Copy extent, effectExtent, graphic from the inline element
            for tag in ("wp:extent", "wp:effectExtent", "a:graphic"):
                child = inline.find(qn(tag))
                if child is not None:
                    anchor.append(child)

            # <wp:wrapNone> — no text wrapping (behind text)
            anchor.append(OxmlElement("wp:wrapNone"))

            # <wp:docPr> — unique id
            docpr = OxmlElement("wp:docPr")
            docpr.set("id",   "9001")
            docpr.set("name", "CompanyWatermark")
            anchor.append(docpr)

            # Replace <wp:inline> with <wp:anchor>
            drawing = inline.getparent()
            if drawing is not None:
                drawing.remove(inline)
                drawing.append(anchor)

    except Exception:
        pass   # watermark is cosmetic — never break the export


def convert_to_docx(md_path: Path, output_dir: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    md_text = md_path.read_text(encoding="utf-8")
    title   = md_path.stem.replace("-", " ").replace("_", " ").title()
    doc     = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(*hex_to_rgb(SLATE))

    def add_rule(color_hex, thickness="12"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    thickness)
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color_hex.lstrip("#"))
        pBdr.append(bot)
        pPr.append(pBdr)

    def strip_inline(text: str):
        parts   = []
        pattern = re.compile(r"(`[^`]+`|\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)")
        last    = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                parts.append((text[last:m.start()], False, False, False))
            tok = m.group(0)
            if   tok.startswith("`"):   parts.append((tok[1:-1],   False, False, True))
            elif tok.startswith("***"): parts.append((tok[3:-3],   True,  True,  False))
            elif tok.startswith("**"):  parts.append((tok[2:-2],   True,  False, False))
            else:                       parts.append((tok[1:-1],   False, True,  False))
            last = m.end()
        if last < len(text):
            parts.append((text[last:], False, False, False))
        return parts

    def add_para(text, color=SLATE, size=10.5, indent=0, prefix="",
                 space_before=2, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Inches(0.25 * indent)
        for chunk, bold, italic, code in strip_inline(prefix + text):
            run = p.add_run(chunk)
            run.bold   = bold
            run.italic = italic
            if code:
                run.font.name  = "Courier New"
                run.font.size  = Pt(9)
                run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))
            else:
                run.font.size  = Pt(size)
                run.font.color.rgb = RGBColor(*hex_to_rgb(color))
        return p

    def flush_table(rows):
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=col_count)
        t.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(col_count):
                cell_text = row[ci].strip() if ci < len(row) else ""
                cell      = t.cell(ri, ci)
                cell.text = cell_text
                para      = cell.paragraphs[0]
                if para.runs:
                    para.runs[0].font.size = Pt(9.5)
                    if ri == 0:
                        para.runs[0].bold = True
                        para.runs[0].font.color.rgb = RGBColor(*hex_to_rgb(NAVY))
                        tc   = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd  = OxmlElement("w:shd")
                        shd.set(qn("w:val"),   "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"),  ICE_BLUE.lstrip("#"))
                        tcPr.append(shd)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Document header ──
    p  = doc.add_paragraph()
    r1 = p.add_run(COMPANY_NAME); r1.bold = True; r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))
    r3 = p.add_run("  —  Product Requirements Document")
    r3.font.size = Pt(9); r3.font.color.rgb = RGBColor(*hex_to_rgb(WARM_TAN))
    p.paragraph_format.space_after = Pt(4)
    add_rule(GOLD, thickness="18")

    # ── Parse pre-processed lines ──
    lines      = preprocess(md_text)
    i          = 0
    in_code    = False
    code_buf   = []
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        m_fence = re.match(r"^(`{3,})\s*(\S*)\s*$", line.rstrip())
        if m_fence and not in_code:
            in_code  = True
            code_buf = []
            i += 1
            continue
        if in_code:
            if re.match(r"^(`{3,})\s*$", line.rstrip()):
                in_code = False
                p = doc.add_paragraph("\n".join(code_buf))
                p.paragraph_format.left_indent  = Inches(0.3)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(8)
                for run in p.runs:
                    run.font.name = "Courier New"; run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  OFF_WHITE.lstrip("#"))
                pPr.append(shd)
            else:
                code_buf.append(line)
            i += 1
            continue

        # Tables
        if "|" in line and line.strip().startswith("|"):
            cells  = [c for c in line.strip().split("|") if c.strip()]
            is_sep = all(re.match(r"[-: ]+$", c.strip()) for c in cells)
            if not is_sep:
                table_rows.append(cells)
            i += 1
            if i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                continue
            flush_table(table_rows)
            table_rows = []
            continue

        if table_rows:
            flush_table(table_rows)
            table_rows = []

        # Headings
        if   line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(*hex_to_rgb(NAVY))
                p.runs[0].font.size = Pt(20)
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
        elif line.startswith("## "):
            add_rule(GRAY, thickness="6")
            p = doc.add_heading(line[3:].strip(), level=2)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(*hex_to_rgb(NAVY))
                p.runs[0].font.size = Pt(13)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(*hex_to_rgb(SLATE))
                p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            if p.runs: p.runs[0].font.size = Pt(10.5)
            p.paragraph_format.space_before = Pt(8)
        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            add_rule(GRAY)
        # Blockquote
        elif line.startswith("> "):
            add_para(line[2:], color=NAVY, size=9.5, indent=1)
        # Checklist
        elif re.match(r"^- \[[ xX]\] ", line):
            add_para(re.sub(r"^- \[[ xX]\] ", "", line), indent=1, prefix="☐  ")
        # Indented list
        elif re.match(r"^  [*\-] ", line):
            add_para(line[4:].strip(), indent=2, prefix="◦  ")
        # Unordered list
        elif re.match(r"^[*\-] ", line):
            add_para(line[2:].strip(), indent=1, prefix="•  ")
        # Ordered list
        elif re.match(r"^\d+\. ", line):
            m = re.match(r"^(\d+)\. (.+)", line)
            if m: add_para(m.group(2), indent=1, prefix=f"{m.group(1)}.  ")
        # Empty line
        elif line.strip() == "":
            pass
        # Normal paragraph
        else:
            add_para(line)

        i += 1

    # ── Footer ──
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_rule(GOLD, thickness="18")
    p = doc.add_paragraph()
    r = p.add_run(f"{COMPANY_NAME} — Confidential  |  Generated by {COMPANY_NAME} Feature Breakdown skill")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(*hex_to_rgb(WARM_TAN))

    _add_docx_watermark(doc)

    out_path = output_dir / (md_path.stem + ".docx")
    doc.save(str(out_path))
    return out_path


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description=f"Convert a {COMPANY_NAME} PRD markdown file to PDF and/or DOCX."
    )
    parser.add_argument("input",       help="Path to the .md PRD file")
    parser.add_argument("--output", "-o", help="Output directory (default: same as input)")
    parser.add_argument("--pdf-only",  action="store_true")
    parser.add_argument("--docx-only", action="store_true")
    args = parser.parse_args()

    md_path = Path(args.input).resolve()
    if not md_path.exists():
        print(f"Error: file not found — {md_path}", file=sys.stderr); sys.exit(1)
    if md_path.suffix.lower() != ".md":
        print(f"Error: expected .md, got {md_path.suffix}", file=sys.stderr); sys.exit(1)

    output_dir = Path(args.output).resolve() if args.output else md_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n{COMPANY_NAME} PRD Converter")
    print(f"Input : {md_path}")
    print(f"Output: {output_dir}\n")

    if not args.docx_only:
        print("Generating PDF...", end=" ", flush=True)
        try:
            p = convert_to_pdf(md_path, output_dir)
            print(f"Done  →  {p.name}")
        except Exception as e:
            print(f"Failed\n{e}", file=sys.stderr); raise

    if not args.pdf_only:
        print("Generating DOCX...", end=" ", flush=True)
        try:
            p = convert_to_docx(md_path, output_dir)
            print(f"Done  →  {p.name}")
        except Exception as e:
            print(f"Failed\n{e}", file=sys.stderr); raise

    print("\nAll done.")


if __name__ == "__main__":
    main()
